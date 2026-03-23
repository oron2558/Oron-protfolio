from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# ==================== TWO-COLUMN TABLE ====================
# Create a 1-row, 2-column table for the layout
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

# Set column widths: left ~65%, right ~35%
page_width = Cm(18)  # 21 - 1.5 - 1.5
left_width = Cm(11.5)
right_width = Cm(6.5)

table.columns[0].width = left_width
table.columns[1].width = right_width

# Remove table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

left_cell = table.cell(0, 0)
right_cell = table.cell(0, 1)

# Set right cell background to light gray
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
right_cell._tc.get_or_add_tcPr().append(shading)

# Set cell padding
for cell in [left_cell, right_cell]:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="120" w:type="dxa"/>'
        '  <w:left w:w="160" w:type="dxa"/>'
        '  <w:bottom w:w="120" w:type="dxa"/>'
        '  <w:right w:w="160" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)

# ==================== HELPER FUNCTIONS ====================

def add_para(cell, text='', bold=False, size=10, color=None, space_after=0, space_before=0, alignment=None):
    p = cell.paragraphs[-1] if cell.paragraphs and cell.paragraphs[-1].text == '' and len(cell.paragraphs) == 1 else cell.add_paragraph()
    # If the cell already has content, always add new paragraph
    if p.text != '' or len(cell.paragraphs) > 1:
        p = cell.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if alignment:
        p.alignment = alignment
    return p

def add_text(cell, runs_data, space_after=0, space_before=0):
    """Add paragraph with multiple runs. runs_data = [(text, size, bold, color), ...]"""
    # Check if we can use existing empty paragraph
    if cell.paragraphs and cell.paragraphs[-1].text == '' and len(cell.paragraphs) == 1:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    for text, size, bold, color in runs_data:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_section_header(cell, title, color=(0x4A, 0x86, 0xC8)):
    """Blue section header like in the reference"""
    p = cell.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_sidebar_header(cell, title, color=(0x4A, 0x86, 0xC8)):
    """Blue section header for sidebar"""
    p = cell.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet_item(cell, text, size=9.5, color=(0x33, 0x33, 0x33)):
    p = cell.add_paragraph()
    run = p.add_run('•  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(10)
    return p

def add_sidebar_bullet(cell, text, size=9.5, color=(0x33, 0x33, 0x33)):
    p = cell.add_paragraph()
    run = p.add_run('•  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(8)
    return p

# ==================== LEFT COLUMN ====================

# Name
p = left_cell.paragraphs[0]
run = p.add_run('ORON TURGEMAN')
run.font.name = 'Calibri'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p.paragraph_format.space_after = Pt(2)

# Role
p = left_cell.add_paragraph()
run = p.add_run('UX DESIGNER \\ PRODUCT DESIGNER')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.bold = False
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(12)

# --- Summary ---
add_section_header(left_cell, 'Summary')

p = left_cell.add_paragraph()
run = p.add_run('A results-driven ')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run = p.add_run('Project Manager and Digital Strategist')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run = p.add_run(' with extensive experience in marketing, advertising, and digital project management. Strong background in market research, strategic planning, and executing digital solutions from concept to launch. Proven ability to collaborate across teams, including development, design, media, and creativity, while managing high-budget campaigns and digital assets.')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.25

p = left_cell.add_paragraph()
run = p.add_run('During my studies, I got very excited about the combination of design and user centered thinking. I welcome you to visit my website and see some of my work. Now I am looking forward to taking it to next level.')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.line_spacing = 1.25

# --- Experience ---
add_section_header(left_cell, 'Experience')

# Job 1 - Africa Israel
p = left_cell.add_paragraph()
run = p.add_run('Jan 2023 – Present')
run.font.name = 'Calibri'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.paragraph_format.space_after = Pt(2)

p = left_cell.add_paragraph()
run = p.add_run('Advertising & Digital Project Coordinator | Africa Israel Residences')
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p.paragraph_format.space_after = Pt(4)

bullets_africa = [
    'Managed and led marketing projects from research and strategy to execution.',
    'Collaborated with development, design, and media teams to create tailored solutions for target audiences.',
    'Developed project briefs and translated business needs into precise marketing and strategic messages.',
    'Managed multimillion-shekel advertising budgets, overseeing vendors and performance monitoring.',
    'Organized corporate events, sales days, and conferences, handling logistics and execution.',
]
for b in bullets_africa:
    add_bullet_item(left_cell, b)

# Job 2 - Hamashbir
p = left_cell.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run('April 2020 – Dec 2022')
run.font.name = 'Calibri'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.paragraph_format.space_after = Pt(2)

p = left_cell.add_paragraph()
run = p.add_run('Advertising & Strategic Project Manager | Hamashbir Lazarchan')
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p.paragraph_format.space_after = Pt(4)

bullets_hamashbir = [
    'Led cross-channel marketing projects from concept development to execution.',
    'Managed and executed marketing strategies, including digital transformation initiatives.',
    'Developed annual marketing plans and guided campaigns from planning to launch.',
    'Worked closely with international suppliers and strategic partners on brand integration.',
    'Planned and managed corporate events, product launches, and branding initiatives.',
]
for b in bullets_hamashbir:
    add_bullet_item(left_cell, b)

# Job 3 - SaaS Founder
p = left_cell.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run('2024 – Present')
run.font.name = 'Calibri'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.paragraph_format.space_after = Pt(2)

p = left_cell.add_paragraph()
run = p.add_run('Founder & Product Designer | SaaS Platform (Self-Employed)')
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p.paragraph_format.space_after = Pt(4)

bullets_saas = [
    'Independently conceived, designed, developed, and launched a full SaaS platform for managing arrival and departure approvals.',
    'Led end-to-end product design from user research to high-fidelity prototypes.',
    'Designed and built the complete UI/UX for web and mobile platforms.',
    'Managed product development lifecycle, shipping a revenue-generating app.',
    'Conducted user testing and iterated based on real-world feedback.',
]
for b in bullets_saas:
    add_bullet_item(left_cell, b)


# ==================== RIGHT COLUMN (SIDEBAR) ====================

# Contact info
p = right_cell.paragraphs[0]
run = p.add_run('📞  052-4541258')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.space_before = Pt(8)

p = right_cell.add_paragraph()
run = p.add_run('✉  oron2558@gmail.com')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
p.paragraph_format.space_after = Pt(4)

p = right_cell.add_paragraph()
run = p.add_run('🌐  www.oronturgemanux.com')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
p.paragraph_format.space_after = Pt(4)

# --- Education ---
add_sidebar_header(right_cell, 'Education')

p = right_cell.add_paragraph()
run = p.add_run('Figma master class & Portfolio')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p.paragraph_format.space_after = Pt(1)

p = right_cell.add_paragraph()
run = p.add_run('Certification Program - by Dan Shiran')
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(1)

p = right_cell.add_paragraph()
run = p.add_run('(2025-2026)')
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(10)

p = right_cell.add_paragraph()
run = p.add_run('User Experience Design')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
p.paragraph_format.space_after = Pt(1)

p = right_cell.add_paragraph()
run = p.add_run('UXV Certification Program - by award winner Tal Florentin')
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(1)

p = right_cell.add_paragraph()
run = p.add_run('(2024-2025)')
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(6)

# --- Skills ---
add_sidebar_header(right_cell, 'Skills')

skills = ['Figma', 'Sketch', 'UX PILOT', 'Loveable', 'Claude AI / Claude Code', 'Google Stitch AI', 'User Research', 'Wireframing & Prototyping', 'Design Systems', 'Product Strategy', 'Project Management', 'Digital Marketing']
for s in skills:
    add_sidebar_bullet(right_cell, s)

# --- Languages ---
add_sidebar_header(right_cell, 'Languages')

p = right_cell.add_paragraph()
run = p.add_run('Hebrew')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
run = p.add_run(' – Native speaker')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(4)

p = right_cell.add_paragraph()
run = p.add_run('English')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
run = p.add_run(' – Fully proficient')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(4)

p = right_cell.add_paragraph()
run = p.add_run('Spanish')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
run = p.add_run(' – Fully proficient')
run.font.name = 'Calibri'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(4)


# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Oron_Turgeman_Resume.docx')
doc.save(output_path)
print(f'Resume saved to: {output_path}')
