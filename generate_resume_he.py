from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# Set RTL for normal style
pPr = style.element.get_or_add_pPr()
bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
pPr.append(bidi)

# ==================== TWO-COLUMN TABLE ====================
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

# Set column widths: right (main) ~65%, left (sidebar) ~35%
# In RTL, first column visually appears on the right
left_width = Cm(6.5)   # sidebar
right_width = Cm(11.5)  # main content

table.columns[0].width = right_width
table.columns[1].width = left_width

# Remove table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

# Make table RTL
bidiVisual = parse_xml(f'<w:bidiVisual {nsdecls("w")}/>')
tblPr.append(bidiVisual)

main_cell = table.cell(0, 0)  # Right side (main content in RTL)
sidebar_cell = table.cell(0, 1)  # Left side (sidebar in RTL)

# Set sidebar background
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
sidebar_cell._tc.get_or_add_tcPr().append(shading)

# Cell padding
for cell in [main_cell, sidebar_cell]:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="120" w:type="dxa"/>'
        '  <w:left w:w="160" w:type="dxa"/>'
        '  <w:bottom w:w="120" w:type="dxa"/>'
        '  <w:right w:w="160" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)

# ==================== HELPERS ====================

def make_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
    pPr.append(bidi)

def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = parse_xml(f'<w:rtl {nsdecls("w")}/>')
    rPr.append(rtl)
    # Set CS font for Hebrew
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial" w:hint="cs"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:cs'), 'Arial')
        rFonts.set(qn('w:hint'), 'cs')

def add_rtl_para(cell, text='', bold=False, size=10, color=None, space_after=0, space_before=0):
    if cell.paragraphs and cell.paragraphs[-1].text == '' and len(cell.paragraphs) == 1:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    make_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        set_rtl_run(run)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_section_header(cell, title, color=(0x4A, 0x86, 0xC8)):
    p = cell.add_paragraph()
    make_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_sidebar_header(cell, title, color=(0x4A, 0x86, 0xC8)):
    p = cell.add_paragraph()
    make_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_bullet(cell, text, size=9.5, color=(0x33, 0x33, 0x33)):
    p = cell.add_paragraph()
    make_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text + '  •')
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.right_indent = Pt(10)

def add_sidebar_bullet(cell, text, size=9.5, color=(0x33, 0x33, 0x33)):
    p = cell.add_paragraph()
    make_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text + '  •')
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Pt(8)

# ==================== MAIN COLUMN (RIGHT IN RTL) ====================

# Name
p = main_cell.paragraphs[0]
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('אורון טורגמן')
run.font.name = 'Arial'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(2)

# Role
p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('מעצב UX/UI \\ מעצב מוצר')
run.font.name = 'Arial'
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(12)

# --- Summary ---
add_section_header(main_cell, 'תקציר')

p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

run = p.add_run('מעצב UX/UI ומעצב מוצר מוכוון תוצאות')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_rtl_run(run)

run = p.add_run(' עם רקע חזק באסטרטגיה דיגיטלית, שיווק וניהול פרויקטים. מנוסה בהובלת מוצרים ממחקר וקונספט ועד השקה — כולל ')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_rtl_run(run)

run = p.add_run('פלטפורמת SaaS פעילה')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_rtl_run(run)

run = p.add_run(' שנבנתה, עוצבה והושקה באופן עצמאי. משלב חשיבה עיצובית מוכוונת משתמש עם הבנה עסקית אמיתית, לאחר ניהול קמפיינים של מיליוני שקלים וצוותים חוצי-ארגון בפיתוח, עיצוב ומדיה.')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.25

p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('במהלך הלימודים שלי, גיליתי את הכוח שבשילוב עיצוב עם חשיבה מוכוונת משתמש. עכשיו אני מחפש לקחת את זה לשלב הבא — להצטרף לצוות מוצר שבו אוכל ליצור חוויות דיגיטליות משמעותיות ובעלות השפעה.')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.line_spacing = 1.25

# --- Experience ---
add_section_header(main_cell, 'ניסיון תעסוקתי')

# Job 1 - Africa Israel
p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('ינואר 2023 – היום')
run.font.name = 'Arial'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(2)

p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('רכז פרסום ופרויקטים דיגיטליים | אפריקה ישראל מגורים')
run.font.name = 'Arial'
run.font.size = Pt(10.5)
run.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

bullets_africa = [
    'ניהול והובלת פרויקטים שיווקיים ממחקר ואסטרטגיה ועד ביצוע',
    'שיתוף פעולה עם צוותי פיתוח, עיצוב ומדיה ליצירת פתרונות מותאמים לקהלי יעד',
    'פיתוח בריפים לפרויקטים ותרגום צרכים עסקיים למסרים שיווקיים ואסטרטגיים מדויקים',
    'ניהול תקציבי פרסום של מיליוני שקלים, פיקוח על ספקים ומעקב ביצועים',
    'ארגון אירועי חברה, ימי מכירות וכנסים, טיפול בלוגיסטיקה וביצוע',
]
for b in bullets_africa:
    add_bullet(main_cell, b)

# Job 2 - Hamashbir
p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('אפריל 2020 – דצמבר 2022')
run.font.name = 'Arial'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
set_rtl_run(run)
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)

p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('מנהל פרסום ופרויקטים אסטרטגיים | המשביר לצרכן')
run.font.name = 'Arial'
run.font.size = Pt(10.5)
run.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

bullets_hamashbir = [
    'הובלת פרויקטים שיווקיים חוצי-ערוצים מפיתוח קונספט ועד ביצוע',
    'ניהול וביצוע אסטרטגיות שיווק, כולל יוזמות טרנספורמציה דיגיטלית',
    'פיתוח תוכניות שיווק שנתיות והנחיית קמפיינים מתכנון ועד השקה',
    'עבודה צמודה עם ספקים בינלאומיים ושותפים אסטרטגיים בשילוב מותגים',
    'תכנון וניהול אירועי חברה, השקות מוצרים ויוזמות מיתוג',
]
for b in bullets_hamashbir:
    add_bullet(main_cell, b)

# Job 3 - SaaS
p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2024 – היום')
run.font.name = 'Arial'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
set_rtl_run(run)
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)

p = main_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('מייסד ומעצב מוצר | פלטפורמת SaaS (עצמאי)')
run.font.name = 'Arial'
run.font.size = Pt(10.5)
run.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

bullets_saas = [
    'הקמה, עיצוב, פיתוח והשקה עצמאית של פלטפורמת SaaS מלאה לניהול אישורי הגעה ויציאה',
    'הובלת עיצוב מוצר מקצה לקצה ממחקר משתמשים ועד אבות טיפוס ברזולוציה גבוהה',
    'עיצוב ובניית ממשק UX/UI מלא לפלטפורמות ווב ומובייל',
    'ניהול מחזור חיי פיתוח מוצר, השקת אפליקציה מניבה הכנסות',
    'ביצוע בדיקות משתמשים ואיטרציה על בסיס משוב מהשטח',
]
for b in bullets_saas:
    add_bullet(main_cell, b)


# ==================== SIDEBAR (LEFT IN RTL) ====================

# Contact
p = sidebar_cell.paragraphs[0]
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('052-4541258  📞')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.space_before = Pt(8)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('oron2558@gmail.com  ✉')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('www.oronturgemanux.com  🌐')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

# --- Education ---
add_sidebar_header(sidebar_cell, 'השכלה')

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Figma master class & Portfolio')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(1)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('תוכנית הסמכה - דן שירן')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(1)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('(2025-2026)')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(10)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('עיצוב חוויית משתמש')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(1)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('תוכנית הסמכה UXV - טל פלורנטין (מעצב UX עטור פרסים)')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(1)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('(2024-2025)')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(6)

# --- Skills ---
add_sidebar_header(sidebar_cell, 'כישורים')

skills = ['Figma', 'Sketch', 'UX PILOT', 'Loveable', 'Claude AI / Claude Code', 'Google Stitch AI', 'מחקר משתמשים', 'וויירפריימים ואבות טיפוס', 'מערכות עיצוב', 'אסטרטגיית מוצר', 'ניהול פרויקטים', 'שיווק דיגיטלי']
for s in skills:
    add_sidebar_bullet(sidebar_cell, s)

# --- Languages ---
add_sidebar_header(sidebar_cell, 'שפות')

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('עברית')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
run = p.add_run(' – שפת אם')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('אנגלית')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
run = p.add_run(' – שליטה מלאה')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

p = sidebar_cell.add_paragraph()
make_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('ספרדית')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
set_rtl_run(run)
run = p.add_run(' – שליטה מלאה')
run.font.name = 'Arial'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_rtl_run(run)
p.paragraph_format.space_after = Pt(4)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Oron_Turgeman_Resume_HE.docx')
doc.save(output_path)
print(f'Resume (Hebrew) saved to: {output_path}')
